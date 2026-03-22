#!/usr/bin/env python3
"""
라이트(퍼블리셔): draft-final.md → NotebookLM_final.docx 변환 스크립트
B5 판형, Pretendard 폰트, 전문 서식 적용
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── 경로 설정 ───
BASE_DIR = "/Users/jaydenkang/Desktop/New Projects/20260321_노트북LM 책쓰기"
DRAFT_PATH = os.path.join(BASE_DIR, "draft-final.md")
OUTPUT_PATH = os.path.join(BASE_DIR, "NotebookLM_final.docx")
FONT_DIR = os.path.join(BASE_DIR, "font")

# ─── 색상 상수 ───
COLOR_DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
COLOR_MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_COVER_BG = "1A3C40"       # 표지 배경: 진한 틸
COLOR_PART_BG = "E6F2F0"        # 부 제목 페이지 배경: 옅은 민트
COLOR_COVER_ACCENT = RGBColor(0x7F, 0xC4, 0xBC)  # 표지 부제/구분선 포인트색
BOX_BG = "E8F0FE"
BOX_BORDER = "1F4E79"

# ─── 폰트 설정 ───
FONT_NAME = "Pretendard"
FONT_FALLBACK = "Malgun Gothic"

def get_font():
    """시스템에 Pretendard가 있으면 사용, 없으면 Malgun Gothic"""
    pretendard_path = os.path.join(FONT_DIR, "Pretendard-Regular.otf")
    if os.path.exists(pretendard_path):
        return FONT_NAME
    return FONT_FALLBACK


def set_font_for_run(run, font_name, size_pt, bold=False, color=None, italic=False):
    """run에 폰트 스타일 적용"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # 한글 폰트 설정
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_shading(paragraph, color_hex):
    """단락 배경색 설정 (전체 줄 배경)"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    pPr.append(shd)


def set_paragraph_negative_indent(paragraph, left_cm, right_cm):
    """음수 인덴트로 단락을 여백 영역까지 확장"""
    pf = paragraph.paragraph_format
    pf.left_indent = Cm(-left_cm)
    pf.right_indent = Cm(-right_cm)
    # first_line_indent를 0으로 설정해서 첫 줄 들여쓰기 방지
    pf.first_line_indent = Cm(0)


def add_full_width_shaded_paragraph(doc, font_name, color_hex, text="", size=11,
                                     text_color=None, bold=False, alignment=None,
                                     before=0, after=0, line_spacing=1.2):
    """여백까지 배경색이 확장되는 전체 너비 단락 생성"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        tc = text_color or COLOR_BLACK
        set_font_for_run(run, font_name, size, bold=bold, color=tc)
    set_paragraph_spacing(p, before=before, after=after, line_spacing=line_spacing)
    set_paragraph_shading(p, color_hex)
    set_paragraph_negative_indent(p, 2.5, 2.5)  # 양쪽 여백 2.5cm만큼 확장
    return p


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.2):
    """문단 간격 설정"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_page_break(doc):
    """페이지 나누기"""
    doc.add_page_break()


def add_box_paragraph(doc, text, font_name, is_title=False):
    """핵심 정리 박스 내 문단 추가"""
    p = doc.add_paragraph()
    # 박스 배경색 설정
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{BOX_BG}" w:val="clear"/>')
    pPr.append(shading)
    # 테두리 설정
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="4" w:color="{BOX_BORDER}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="4" w:color="{BOX_BORDER}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{BOX_BORDER}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="4" w:color="{BOX_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    # 내부 여백
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="240" w:right="240"/>')
    pPr.append(ind)

    if is_title:
        clean = strip_markdown(text)
        run = p.add_run(clean)
        set_font_for_run(run, font_name, 12, bold=True, color=COLOR_DARK_BLUE)
    else:
        process_inline_formatting(p, text, font_name, 11, COLOR_BLACK)

    set_paragraph_spacing(p, before=4, after=4, line_spacing=1.2)
    return p


def add_styled_table(doc, headers, rows, font_name):
    """스타일이 적용된 표 생성"""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 헤더 행
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        clean_header = strip_markdown(header_text.strip())
        run = p.add_run(clean_header)
        set_font_for_run(run, font_name, 10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 배경색
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 데이터 행
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                # 셀 안의 마크다운 서식을 실제 서식으로 변환
                process_inline_formatting(p, cell_text.strip(), font_name, 10, COLOR_BLACK)
                # 줄무늬 배경
                if row_idx % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F7FB" w:val="clear"/>')
                    cell._element.get_or_add_tcPr().append(shading)

    # 테이블 테두리
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="CCCCCC"/>'
        f'<w:left w:val="single" w:sz="4" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/>'
        f'<w:right w:val="single" w:sz="4" w:color="CCCCCC"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="CCCCCC"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # 표 후 간격
    return table


def parse_markdown_table(lines, start_idx):
    """마크다운 표를 파싱하여 헤더와 행 데이터 반환"""
    headers = []
    rows = []
    idx = start_idx

    if idx < len(lines) and '|' in lines[idx]:
        # 헤더 행
        header_line = lines[idx].strip().strip('|')
        headers = [h.strip() for h in header_line.split('|')]
        idx += 1

        # 구분선 행 (---|---)
        if idx < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[idx]):
            idx += 1

        # 데이터 행
        while idx < len(lines) and '|' in lines[idx] and lines[idx].strip().startswith('|'):
            data_line = lines[idx].strip().strip('|')
            row = [c.strip() for c in data_line.split('|')]
            rows.append(row)
            idx += 1

    return headers, rows, idx


def add_caption(doc, text, font_name):
    """이미지 캡션 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font_for_run(run, font_name, 10, italic=True, color=COLOR_GRAY)
    set_paragraph_spacing(p, before=2, after=8)
    return p


def add_image_placeholder(doc, description, font_name):
    """전문적인 이미지 플레이스홀더: 실선 테두리 + 아이콘 + 설명"""
    # 상단 간격
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=4, after=0)

    # 플레이스홀더 박스 (3줄: 아이콘 행, 설명 행, 안내 행)
    # --- 아이콘 행 ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr1 = p1._element.get_or_add_pPr()
    shading1 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
    pPr1.append(shading1)
    pBdr1 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'<w:left w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'<w:right w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'</w:pBdr>'
    )
    pPr1.append(pBdr1)
    run1 = p1.add_run("IMAGE")
    set_font_for_run(run1, font_name, 14, bold=True, color=COLOR_MED_BLUE)
    set_paragraph_spacing(p1, before=0, after=0, line_spacing=1.2)

    # --- 설명 행 ---
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p2._element.get_or_add_pPr()
    shading2 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
    pPr2.append(shading2)
    pBdr2 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'<w:right w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'</w:pBdr>'
    )
    pPr2.append(pBdr2)
    # 설명에서 [그림 N: 을 제거하고 순수 설명만 추출
    clean_desc = re.sub(r'^\[그림\s*\d*\s*:?\s*', '', description).rstrip(']')
    run2 = p2.add_run(clean_desc)
    set_font_for_run(run2, font_name, 10, color=RGBColor(0x33, 0x33, 0x33))
    set_paragraph_spacing(p2, before=0, after=0, line_spacing=1.2)

    # --- 안내 행 ---
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr3 = p3._element.get_or_add_pPr()
    shading3 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
    pPr3.append(shading3)
    pBdr3 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'<w:right w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="12" w:color="2E75B6"/>'
        f'</w:pBdr>'
    )
    pPr3.append(pBdr3)
    run3 = p3.add_run("(이 위치에 해당 이미지를 삽입하세요)")
    set_font_for_run(run3, font_name, 9, italic=True, color=COLOR_GRAY)
    set_paragraph_spacing(p3, before=0, after=0, line_spacing=1.2)

    # 캡션 (그림 번호)
    fig_match = re.match(r'\[(그림\s*\d+)', description)
    if fig_match:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"{fig_match.group(1)}. {clean_desc}")
        set_font_for_run(cap_run, font_name, 10, italic=False, color=COLOR_GRAY)
        set_paragraph_spacing(caption, before=2, after=8)

    return p2


def strip_markdown(text):
    """마크다운 문법 기호를 모두 제거하고 순수 텍스트만 반환"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **볼드** → 볼드
    text = re.sub(r'\*(.*?)\*', r'\1', text)       # *이탤릭* → 이탤릭
    text = re.sub(r'`(.*?)`', r'\1', text)         # `코드` → 코드
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text) # [링크](url) → 링크
    return text


def process_inline_formatting(paragraph, text, font_name, size=11, color=COLOR_BLACK, base_bold=False):
    """인라인 마크다운(**볼드**, *이탤릭*, `코드`) 처리. ** 기호를 제거하고 실제 서식으로 변환"""
    # **볼드**, *이탤릭*, `코드`를 모두 처리
    pattern = r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            clean = part[2:-2]
            if clean:
                run = paragraph.add_run(clean)
                set_font_for_run(run, font_name, size, bold=True, color=color)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            clean = part[1:-1]
            if clean:
                run = paragraph.add_run(clean)
                set_font_for_run(run, font_name, size, bold=base_bold, color=color, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            clean = part[1:-1]
            if clean:
                run = paragraph.add_run(clean)
                set_font_for_run(run, font_name, size, bold=base_bold, color=RGBColor(0x88, 0x33, 0x00))
        else:
            run = paragraph.add_run(part)
            set_font_for_run(run, font_name, size, bold=base_bold, color=color)


def create_cover_page(doc, font_name):
    """표지 페이지 생성 - 다크 틸 배경 (전체 페이지 채움)"""
    bg = COLOR_COVER_BG
    center = WD_ALIGN_PARAGRAPH.CENTER

    # 상단 여백 채우기
    for _ in range(10):
        add_full_width_shaded_paragraph(doc, font_name, bg, before=0, after=0)

    # 제목
    add_full_width_shaded_paragraph(doc, font_name, bg,
        text="노트북LM으로 다 됨", size=28, text_color=COLOR_WHITE,
        bold=True, alignment=center, before=0, after=8)

    # 구분선
    add_full_width_shaded_paragraph(doc, font_name, bg,
        text="━━━━━━━━━━━━━━━", size=10, text_color=COLOR_COVER_ACCENT,
        alignment=center, before=4, after=4)

    # 부제
    add_full_width_shaded_paragraph(doc, font_name, bg,
        text="팀장을 위한 실전 활용법 20선", size=16, text_color=COLOR_COVER_ACCENT,
        alignment=center, before=0, after=24)

    # 저자
    add_full_width_shaded_paragraph(doc, font_name, bg,
        text="AI ROASTING", size=14, text_color=COLOR_WHITE,
        alignment=center, before=0, after=12)

    # 작성일
    today = datetime.date.today().strftime("%Y년 %m월 %d일")
    add_full_width_shaded_paragraph(doc, font_name, bg,
        text=today, size=12, text_color=RGBColor(0x99, 0xBB, 0xB7),
        alignment=center, before=0, after=0)

    # 하단 여백 채우기
    for _ in range(8):
        add_full_width_shaded_paragraph(doc, font_name, bg, before=0, after=0)

    add_page_break(doc)


def create_toc(doc, font_name):
    """목차 페이지 생성"""
    p = doc.add_paragraph()
    run = p.add_run("목차")
    set_font_for_run(run, font_name, 20, bold=True, color=COLOR_DARK_BLUE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=24)

    # TOC 필드 코드 삽입
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)
    run4 = p.add_run("(목차를 업데이트하려면 이 영역을 우클릭하고 '필드 업데이트'를 선택하세요)")
    set_font_for_run(run4, font_name, 10, color=COLOR_GRAY)
    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)

    add_page_break(doc)


def setup_document(doc, font_name):
    """문서 기본 설정: B5 판형, 여백, 헤더/푸터"""
    # B5 판형 설정
    section = doc.sections[0]
    section.page_width = Cm(17.6)   # 176mm
    section.page_height = Cm(25.0)  # 250mm
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 헤더
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("노트북LM으로 다 됨")
    set_font_for_run(run, font_name, 9, color=COLOR_GRAY)
    # 헤더 하단 선
    pPr = hp._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 푸터 (페이지 번호)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_font_for_run(run, font_name, 9, color=COLOR_GRAY)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)
    run4 = fp.add_run("1")
    set_font_for_run(run4, font_name, 9, color=COLOR_GRAY)
    run5 = fp.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)

    # 문서 속성
    doc.core_properties.author = "AI ROASTING"
    doc.core_properties.title = "노트북LM으로 다 됨: 팀장을 위한 실전 활용법 20선"


def process_markdown(doc, md_text, font_name):
    """마크다운 텍스트를 Word 문서로 변환"""
    lines = md_text.split('\n')
    i = 0
    in_box = False
    in_quote = False
    quote_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 주석 처리
        if stripped.startswith('<!-- pagebreak -->'):
            add_page_break(doc)
            i += 1
            continue

        if stripped == '<!-- box-start -->':
            in_box = True
            i += 1
            continue

        if stripped == '<!-- box-end -->':
            in_box = False
            i += 1
            continue

        # 빈 줄
        if not stripped:
            if in_quote and quote_buffer:
                # 인용 블록 종료
                quote_text = '\n'.join(quote_buffer)
                p = doc.add_paragraph()
                pPr = p._element.get_or_add_pPr()
                # 왼쪽 테두리로 인용 표시
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:left w:val="single" w:sz="12" w:space="8" w:color="2E75B6"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
                ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="480"/>')
                pPr.append(ind)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA" w:val="clear"/>')
                pPr.append(shading)
                process_inline_formatting(p, quote_text, font_name, 10.5, RGBColor(0x33, 0x33, 0x33))
                set_paragraph_spacing(p, before=6, after=6, line_spacing=1.2)
                quote_buffer = []
                in_quote = False
            i += 1
            continue

        # 인용 블록 (> 로 시작)
        if stripped.startswith('>'):
            in_quote = True
            quote_content = stripped[1:].strip()
            quote_buffer.append(quote_content)
            i += 1
            continue
        elif in_quote and quote_buffer:
            # 인용 블록이 끝남
            quote_text = '\n'.join(quote_buffer)
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="12" w:space="8" w:color="2E75B6"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="480"/>')
            pPr.append(ind)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA" w:val="clear"/>')
            pPr.append(shading)
            run = p.add_run(quote_text)
            set_font_for_run(run, font_name, 10.5, italic=True, color=RGBColor(0x33, 0x33, 0x33))
            set_paragraph_spacing(p, before=6, after=6, line_spacing=1.2)
            quote_buffer = []
            in_quote = False

        # 구분선
        if stripped == '---':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            set_paragraph_spacing(p, before=6, after=6)
            i += 1
            continue

        # 표 감지
        if '|' in stripped and stripped.startswith('|'):
            headers, rows, new_idx = parse_markdown_table(lines, i)
            if headers and rows:
                add_styled_table(doc, headers, rows, font_name)
            i = new_idx
            continue

        # [그림 N: ...] 이미지 삽입 또는 플레이스홀더
        img_match = re.match(r'\[그림\s*(\d+)\s*[:.]\s*(.*?)\]', stripped)
        if img_match:
            fig_num = int(img_match.group(1))
            fig_desc = img_match.group(2).strip()
            img_path = os.path.join(BASE_DIR, "images", f"fig{fig_num:02d}.png")
            if os.path.exists(img_path):
                # 실제 이미지 삽입
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # B5 기준 본문 폭 약 12.6cm, 이미지는 그보다 약간 작게
                run.add_picture(img_path, width=Cm(11.5))
                set_paragraph_spacing(p, before=8, after=2)
                # 캡션
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(f"그림 {fig_num}. {fig_desc}")
                set_font_for_run(cap_run, font_name, 10, color=COLOR_GRAY)
                set_paragraph_spacing(cap_p, before=2, after=10)
            else:
                add_image_placeholder(doc, img_match.group(0), font_name)
            i += 1
            continue
        # [그림: ...] 번호 없는 이미지 (이미지 없이 스킵)
        img_match2 = re.match(r'\[그림\s*:?\s*(.*?)\]', stripped)
        if img_match2:
            # 번호 없는 그림 참조는 부연 캡션이므로 건너뛴다
            i += 1
            continue

        # 제목 레벨 처리
        if stripped.startswith('#'):
            level = len(stripped.split()[0])  # # 개수
            title_text = stripped.lstrip('#').strip()

            if level == 1:
                # 부 제목 (# 제1부: ...) - 옅은 민트 배경 전체 페이지
                bg = COLOR_PART_BG
                center = WD_ALIGN_PARAGRAPH.CENTER

                # 페이지 브레이크 (이전 단락에 pageBreakBefore 설정)
                first_p = add_full_width_shaded_paragraph(doc, font_name, bg, before=0, after=0)
                first_p.paragraph_format.page_break_before = True

                # 상단 여백
                for _ in range(7):
                    add_full_width_shaded_paragraph(doc, font_name, bg, before=0, after=0)

                # 부 제목 (Heading 1 스타일 유지하되 배경 확장)
                p = doc.add_paragraph()
                p.style = doc.styles['Heading 1']
                p.clear()
                p.alignment = center
                process_inline_formatting(p, title_text, font_name, 22, COLOR_DARK_BLUE, True)
                set_paragraph_spacing(p, before=0, after=8)
                set_paragraph_shading(p, bg)
                set_paragraph_negative_indent(p, 2.5, 2.5)

                # 구분선
                add_full_width_shaded_paragraph(doc, font_name, bg,
                    text="━━━━━━━━━━", size=10,
                    text_color=RGBColor(0x1A, 0x3C, 0x40),
                    alignment=center, before=4, after=4)

                # 부 제목 다음에 부제가 있으면 같은 페이지에 렌더링
                next_idx = i + 1
                while next_idx < len(lines) and lines[next_idx].strip() == '':
                    next_idx += 1
                if next_idx < len(lines):
                    subtitle_line = lines[next_idx].strip()
                    if subtitle_line.startswith('**') and subtitle_line.endswith('**'):
                        subtitle_text = subtitle_line.strip('*').strip()
                        add_full_width_shaded_paragraph(doc, font_name, bg,
                            text=subtitle_text, size=14,
                            text_color=RGBColor(0x1A, 0x3C, 0x40),
                            alignment=center, before=8, after=0)

                        # 하단 채우기
                        for _ in range(14):
                            add_full_width_shaded_paragraph(doc, font_name, bg, before=0, after=0)

                        i = next_idx + 1
                        continue

                # 부제가 없는 경우 하단 채우기
                for _ in range(14):
                    add_full_width_shaded_paragraph(doc, font_name, bg, before=0, after=0)
                i += 1
                continue

            elif level == 2:
                # 장/활용법 제목 (## 제1장: ...)
                add_page_break(doc)
                p = doc.add_paragraph()
                p.style = doc.styles['Heading 2']
                p.clear()
                process_inline_formatting(p, title_text, font_name, 18, COLOR_DARK_BLUE, True)
                set_paragraph_spacing(p, before=24, after=8)

            elif level == 3:
                # 섹션 제목 (### ...)
                p = doc.add_paragraph()
                p.style = doc.styles['Heading 3']
                p.clear()
                process_inline_formatting(p, title_text, font_name, 14, COLOR_DARK_BLUE, True)
                set_paragraph_spacing(p, before=18, after=6)

            elif level == 4:
                # 하위 섹션 (#### ...)
                p = doc.add_paragraph()
                p.style = doc.styles['Heading 4'] if 'Heading 4' in [s.name for s in doc.styles] else doc.styles['Heading 3']
                p.clear()
                process_inline_formatting(p, title_text, font_name, 12, COLOR_MED_BLUE, True)
                set_paragraph_spacing(p, before=12, after=4)

            i += 1
            continue

        # 리스트 항목
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', stripped)
        if list_match:
            indent_level = len(line) - len(line.lstrip())
            bullet_or_num = list_match.group(2)
            content = list_match.group(3)

            if in_box:
                add_box_paragraph(doc, f"  {'  ' * (indent_level // 2)}{bullet_or_num} {content}", font_name)
            else:
                p = doc.add_paragraph()
                indent_cm = 0.5 + (indent_level // 2) * 0.5
                p.paragraph_format.left_indent = Cm(indent_cm)

                if bullet_or_num in ['-', '*']:
                    prefix = "• "
                else:
                    prefix = f"{bullet_or_num} "

                process_inline_formatting(p, f"{prefix}{content}", font_name, 11, COLOR_BLACK)
                set_paragraph_spacing(p, before=0, after=6, line_spacing=1.2)

            i += 1
            continue

        # 일반 본문
        if in_box:
            # 핵심 정리 내용인지 확인
            if '핵심 정리' in stripped or '✅' in stripped or '📌' in stripped:
                add_box_paragraph(doc, stripped, font_name, is_title=True)
            else:
                add_box_paragraph(doc, stripped, font_name)
        else:
            p = doc.add_paragraph()
            process_inline_formatting(p, stripped, font_name, 11, COLOR_BLACK)
            set_paragraph_spacing(p, before=0, after=6, line_spacing=1.2)

        i += 1

    # 남은 인용 버퍼 처리
    if quote_buffer:
        quote_text = '\n'.join(quote_buffer)
        p = doc.add_paragraph()
        run = p.add_run(quote_text)
        set_font_for_run(run, font_name, 10.5, italic=True, color=RGBColor(0x33, 0x33, 0x33))


def main():
    print("=" * 60)
    print("라이트(퍼블리셔): Word 문서 생성 시작")
    print("=" * 60)

    # 폰트 확인
    font_name = get_font()
    print(f"사용 폰트: {font_name}")

    # 마크다운 읽기
    print("draft-final.md 읽는 중...")
    with open(DRAFT_PATH, 'r', encoding='utf-8') as f:
        md_text = f.read()
    print(f"총 {len(md_text):,}자 읽음")

    # 문서 생성
    doc = Document()
    print("문서 기본 설정 중...")
    setup_document(doc, font_name)

    # 표지
    print("표지 생성 중...")
    create_cover_page(doc, font_name)

    # 목차
    print("목차 생성 중...")
    create_toc(doc, font_name)

    # 본문 변환
    print("본문 변환 중... (시간이 걸릴 수 있습니다)")

    # 첫 줄 (제목, 저자)은 표지에서 처리했으므로 건너뛰기
    # "# 노트북LM으로 다 됨"과 "**AI ROASTING 지음**" 줄 건너뛰기
    lines = md_text.split('\n')
    skip_idx = 0
    for idx, line in enumerate(lines):
        if line.strip().startswith('# 노트북LM') or line.strip().startswith('**AI ROASTING'):
            skip_idx = idx + 1
        elif line.strip() and skip_idx > 0:
            break

    # 건너뛴 후의 본문만 처리
    body_text = '\n'.join(lines[skip_idx:])
    process_markdown(doc, body_text, font_name)

    # 저장
    print(f"저장 중: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("=" * 60)
    print(f"완료! {OUTPUT_PATH}")
    print("=" * 60)


if __name__ == '__main__':
    main()
